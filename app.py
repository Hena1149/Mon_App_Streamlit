import sys
import numpy as np

def verify_environment():
    required_python = (3, 9)
    current_python = sys.version_info[:2]
    
    if current_python != required_python:
        raise RuntimeError(f"Python 3.9 required (current: {sys.version})")
    
    required_numpy = "1.21.2"
    if np.__version__ != required_numpy:
        raise RuntimeError(f"numpy {required_numpy} required (current: {np.__version__})")

verify_environment()
import streamlit as st
import os
import tempfile
from utils.file_utils import process_uploaded_file
from utils.text_processing import generate_wordcloud, clean_text
from utils.openai_utils import generate_rules, generate_checkpoints, generate_test_cases
import pyperclip  # Pour la copie dans le presse-papier
from collections import Counter
from docx import Document  # Ajoutez cette ligne avec les autres imports

# Configuration de la page
st.set_page_config(
    page_title="Analyse de Cahier des Charges",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Style CSS personnalisé
st.markdown("""
    <style>
    .main {
        background-color: #f8f9fa;
    }
    .stButton>button {
        background-color: #4CAF50;
        color: white;
        border-radius: 5px;
        padding: 0.5rem 1rem;
    }
    .stFileUploader>div>div>button {
        background-color: #2196F3;
        color: white;
    }
    .sidebar .sidebar-content {
        background-color: #e3f2fd;
    }
    h1 {
        color: #2c3e50;
    }
    </style>
    """, unsafe_allow_html=True)

def main():
    st.title("📄 Analyse Automatisée de Cahier des Charges")
    st.markdown("""
    Chargez votre document (PDF ou Word) pour en extraire :
    - Les règles de gestion
    - Les points de contrôle
    - Les cas de test
    """)

    # Initialisation des variables de session
    if 'text' not in st.session_state:
        st.session_state.text = ""
    if 'rules' not in st.session_state:
        st.session_state.rules = []
    if 'checkpoints' not in st.session_state:
        st.session_state.checkpoints = []
    if 'test_cases' not in st.session_state:
        st.session_state.test_cases = []

    # Sidebar pour les paramètres
    with st.sidebar:
        st.header("Paramètres")
        st.session_state.openai_key = st.text_input("Clé API OpenAI", type="password")
        st.session_state.openai_endpoint = st.text_input("Endpoint Azure OpenAI", "https://chat-genai.openai.azure.com/")
        st.session_state.model_name = st.selectbox("Modèle", ["gpt-4o", "gpt-35-turbo"])
        
        st.divider()
        st.info("Configurez votre clé API et endpoint avant de commencer.")

    # Onglets principaux
    tab1, tab2, tab3, tab4 = st.tabs(["📤 Upload", "📊 Analyse", "✅ Points de contrôle", "🧪 Cas de test"])

    with tab1:
        st.header("Chargement du document")
        uploaded_file = st.file_uploader("Téléversez votre cahier des charges", type=["pdf", "docx"])
        
        if uploaded_file is not None:
            with st.spinner("Extraction du texte en cours..."):
                # Sauvegarde temporaire du fichier
                with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(uploaded_file.name)[1]) as tmp_file:
                    tmp_file.write(uploaded_file.getvalue())
                    tmp_path = tmp_file.name
                
                # Traitement du fichier
                st.session_state.text = process_uploaded_file(tmp_path)
                os.unlink(tmp_path)  # Nettoyage du fichier temporaire
            
            st.success("Texte extrait avec succès !")
            st.expander("Aperçu du texte extrait").text(st.session_state.text[:2000] + "...")

    with tab2:
        st.header("Analyse Textuelle")
        if st.session_state.text:
            col1, col2 = st.columns(2)
            
            with col1:
                st.subheader("Nuage de mots clés")
                with st.spinner("Génération du wordcloud..."):
                    fig = generate_wordcloud(st.session_state.text)
                    st.pyplot(fig)
            
            with col2:
                st.subheader("Mots les plus fréquents")
                tokens = clean_text(st.session_state.text)
                freq_dist = Counter(tokens)
                top_words = freq_dist.most_common(10)
                
                for word, freq in top_words:
                    st.markdown(f"- **{word}**: {freq} occurrences")
                
                st.download_button(
                    label="Télécharger l'analyse",
                    data="\n".join([f"{w}: {f}" for w, f in top_words]),
                    file_name="frequence_mots.txt"
                )
            
            st.divider()
            if st.button("Générer les règles de gestion", disabled=not st.session_state.text):
                with st.spinner("Génération des règles en cours..."):
                    st.session_state.rules = generate_rules(
                        st.session_state.text,
                        st.session_state.openai_key,
                        st.session_state.openai_endpoint,
                        st.session_state.model_name
                    )
                
                if st.session_state.rules:
                    st.success(f"{len(st.session_state.rules)} règles générées !")
                    st.download_button(
                        label="Télécharger les règles",
                        data="\n".join(st.session_state.rules),
                        file_name="regles_gestion.txt"
                    )
        else:
            st.warning("Veuillez d'abord charger un document dans l'onglet Upload.")

    with tab3:
        st.header("Points de Contrôle")
        
        # Section pour l'upload de points existants
        with st.expander("⚙️ Points de contrôle existants (facultatif)"):
            existing_cp_file = st.file_uploader(
                "Uploader un fichier avec des points de contrôle existants (PDF, Word ou TXT)",
                type=["pdf", "docx", "txt"],
                key="existing_cp"
            )
            
            if existing_cp_file:
                with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(existing_cp_file.name)[1]) as tmp_file:
                    tmp_file.write(existing_cp_file.getvalue())
                    tmp_path = tmp_file.name
                
                if existing_cp_file.name.endswith(".txt"):
                    with open(tmp_path, "r", encoding="utf-8") as f:
                        st.session_state.existing_checkpoints = [line.strip() for line in f.readlines() if line.strip()]
                else:
                    extracted_text = process_uploaded_file(tmp_path)
                    st.session_state.existing_checkpoints = [line.strip() for line in extracted_text.split("\n") if line.strip()]
                
                os.unlink(tmp_path)
                st.success(f"{len(st.session_state.existing_checkpoints)} points de contrôle existants chargés !")
                
                # Afficher un aperçu des points existants
                with st.expander("👀 Aperçu des points existants"):
                    st.write(f"Premiers {min(5, len(st.session_state.existing_checkpoints))} points :")
                    for i, point in enumerate(st.session_state.existing_checkpoints[:5]):
                        st.markdown(f"{i+1}. {point}")
        
        if st.session_state.rules:
            if st.button("Générer les points de contrôle"):
                with st.spinner("Création des points de contrôle..."):
                    raw_checkpoints = generate_checkpoints(
                        st.session_state.rules,
                        st.session_state.openai_key,
                        st.session_state.openai_endpoint,
                        st.session_state.model_name
                    )
                    
                    # Suppression des doublons si des points existants ont été uploadés
                    if hasattr(st.session_state, 'existing_checkpoints'):
                        st.session_state.checkpoints = remove_duplicates(
                            raw_checkpoints,
                            st.session_state.existing_checkpoints
                        )
                        duplicates_removed = len(raw_checkpoints) - len(st.session_state.checkpoints)
                        st.info(f"{duplicates_removed} doublons supprimés")
                        
                        # Option pour voir les doublons
                        if duplicates_removed > 0:
                            if st.checkbox("Afficher les doublons détectés", key="show_duplicates"):
                                duplicates = [
                                    point for point in raw_checkpoints
                                    if any(is_similar(point, existing) for existing in st.session_state.existing_checkpoints)
                                ]
                                st.warning("Points considérés comme doublons :")
                                for dup in duplicates[:10]:  # Limiter à 10 pour ne pas surcharger
                                    st.markdown(f"- {dup}")
                                if len(duplicates) > 10:
                                    st.write(f"...et {len(duplicates)-10} autres")
                    else:
                        st.session_state.checkpoints = raw_checkpoints
                    
                    if st.session_state.checkpoints:
                        st.success(f"{len(st.session_state.checkpoints)} points générés !")
                        
                        # Affichage avec pagination
                        st.subheader("Points de contrôle générés")
                        
                        # Sélection du format d'affichage
                        display_format = st.radio(
                            "Format d'affichage",
                            ["Liste paginée", "Liste complète"],
                            horizontal=True,
                            key="cp_display_format"
                        )
                        
                        if display_format == "Liste paginée":
                            page_size = st.slider(
                                "Points par page",
                                min_value=5,
                                max_value=50,
                                value=10,
                                key="cp_page_size"
                            )
                            total_pages = max(1, (len(st.session_state.checkpoints) + page_size - 1) // page_size)
                            page = st.number_input(
                                "Page",
                                min_value=1,
                                max_value=total_pages,
                                value=1,
                                key="cp_page_number"
                            )
                            
                            start_idx = (page - 1) * page_size
                            end_idx = start_idx + page_size
                            
                            for i, cp in enumerate(st.session_state.checkpoints[start_idx:end_idx], start=start_idx+1):
                                st.markdown(f"{i}. {cp}")
                            
                            st.caption(f"Page {page}/{total_pages} | {len(st.session_state.checkpoints)} points au total")
                        else:
                            for i, cp in enumerate(st.session_state.checkpoints, start=1):
                                st.markdown(f"{i}. {cp}")
                        
                        # Options de téléchargement
                        st.divider()
                        st.subheader("Export des résultats")
                        
                        col1, col2 = st.columns(2)
                        
                        with col1:
                            # Format Word
                            doc = Document()
                            doc.add_heading('Points de Contrôle', level=1)
                            for cp in st.session_state.checkpoints:
                                doc.add_paragraph(cp)
                            
                            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                                doc.save(tmp.name)
                                with open(tmp.name, "rb") as f:
                                    st.download_button(
                                        label="📝 Télécharger en Word",
                                        data=f,
                                        file_name="points_de_controle.docx",
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                    )
                            os.unlink(tmp.name)
                        
                        with col2:
                            # Format Texte
                            text_content = "\n".join(f"{i+1}. {cp}" for i, cp in enumerate(st.session_state.checkpoints))
                            st.download_button(
                                label="📄 Télécharger en TXT",
                                data=text_content,
                                file_name="points_de_controle.txt",
                                mime="text/plain"
                            )
                        
                        # Option pour copier dans le clipboard
                        if st.button("📋 Copier dans le presse-papier"):
                            text_to_copy = "\n".join(st.session_state.checkpoints)
                            pyperclip.copy(text_to_copy)
                            st.toast("Points de contrôle copiés dans le presse-papier !", icon="✅")
            else:
                if hasattr(st.session_state, 'checkpoints') and st.session_state.checkpoints:
                    st.warning("Des points de contrôle ont déjà été générés. Cliquez à nouveau pour regénérer.")
        else:
            st.warning("Veuillez d'abord générer les règles de gestion dans l'onglet 'Analyse'.")

    with tab4:
        st.header("Cas de Test")
        if st.session_state.checkpoints:
            if st.button("Générer les cas de test"):
                with st.spinner("Création des cas de test..."):
                    st.session_state.test_cases = generate_test_cases(
                        st.session_state.checkpoints,
                        st.session_state.openai_key,
                        st.session_state.openai_endpoint,
                        st.session_state.model_name
                    )
                
                if st.session_state.test_cases:
                    st.success(f"{len(st.session_state.test_cases)} cas de test générés !")
                    
                    # Sélection d'un cas à afficher
                    selected_case = st.selectbox(
                        "Sélectionnez un cas à visualiser",
                        range(len(st.session_state.test_cases)),
                        format_func=lambda x: f"Cas de test #{x+1}"
                    )
                    
                    st.markdown(st.session_state.test_cases[selected_case])
                    
                    st.download_button(
                        label="Télécharger tous les cas",
                        data="\n".join(st.session_state.test_cases),
                        file_name="cas_de_test.txt"
                    )
        else:
            st.warning("Générez d'abord les points de contrôle dans l'onglet précédent.")

if __name__ == "__main__":
    main()
